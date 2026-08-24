import os
import sys
import sqlite3
import shutil
import uuid
from datetime import datetime, timedelta
import re
import traceback
import subprocess
import threading
import time

from flask import Flask, render_template_string, request, redirect, url_for, flash, session, send_file
from werkzeug.utils import secure_filename

# ============================================================
# تنظیمات لایسنس - برای تست تغییر دهید
# ============================================================
LICENSE_EXPIRY_DAYS = 1  # <--- تغییر این عدد برای تست (مثلاً 1 برای ۱ روز، 7 برای ۷ روز)

# ============================================================
# مسیرها
# ============================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
REPORTS_DIR = os.path.join(BASE_DIR, 'reports')
IMAGES_DIR = os.path.join(BASE_DIR, 'Images')
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'static', 'uploads')

for folder in [REPORTS_DIR, IMAGES_DIR, UPLOAD_FOLDER]:
    if not os.path.exists(folder):
        try:
            os.makedirs(folder)
        except:
            pass

app = Flask(__name__)
app.secret_key = 'mhb_secret_key_2026_fixed'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# ============================================================
# توابع کمکی
# ============================================================
def fa(text):
    if text is None:
        return ""
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display
        reshaped = arabic_reshaper.reshape(str(text))
        return get_display(reshaped)
    except:
        return str(text)

def convert_to_shamsi(miladi_date):
    try:
        if not miladi_date:
            return ""
        if '-' in miladi_date:
            dt = datetime.strptime(miladi_date, "%Y-%m-%d")
        else:
            dt = datetime.strptime(miladi_date.replace('/', '-'), "%Y-%m-%d")
        import jdatetime
        shamsi = jdatetime.date.fromgregorian(date=dt)
        return f"{shamsi.year}/{shamsi.month:02d}/{shamsi.day:02d}"
    except:
        return miladi_date

@app.context_processor
def utility_processor():
    return dict(fa=fa, convert_to_shamsi=convert_to_shamsi)

# ============================================================
# تابع بررسی کتابخانه‌ها (مهم - قبلاً جا افتاده بود)
# ============================================================
def check_report_libs():
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from docx import Document
        return True, "کتابخانه‌ها سالم هستند."
    except ImportError as e:
        return False, f"خطای بارگذاری: {str(e)}"
    except Exception as e:
        return False, f"خطای غیرمنتظره: {str(e)}"

def split_text_to_lines(text, pdf, font_name, max_width, font_size):
    if not text:
        return []
    lines = []
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            lines.append('')
            continue
        pdf.setFont(font_name, font_size)
        text_fa = fa(line)
        if pdf.stringWidth(text_fa, font_name, font_size) <= max_width:
            lines.append(line)
        else:
            words = line.split()
            current_line = ""
            for word in words:
                test_line = current_line + " " + word if current_line else word
                test_line_fa = fa(test_line)
                if pdf.stringWidth(test_line_fa, font_name, font_size) <= max_width:
                    current_line = test_line
                else:
                    if current_line:
                        lines.append(current_line)
                    current_line = word
            if current_line:
                lines.append(current_line)
    return lines

# ============================================================
# دیتابیس
# ============================================================
DB_NAME = "database.db"

def get_db():
    conn = sqlite3.connect(DB_NAME)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db()
    cur = conn.cursor()
    cur.execute("""
    CREATE TABLE IF NOT EXISTS settings(key TEXT PRIMARY KEY, value TEXT)
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS expert(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fullname TEXT, position TEXT, license TEXT,
        phone TEXT, email TEXT, logo_path TEXT,
        signature_path TEXT, stamp_path TEXT)
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS project(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        report_no TEXT, project_name TEXT, employer TEXT,
        address TEXT, visit_date TEXT, stage TEXT,
        description TEXT, logo_path TEXT)
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS defect(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        project_id INTEGER, report_no TEXT, project_name TEXT,
        employer TEXT, visit_date TEXT, stage TEXT,
        title TEXT, standard TEXT, description TEXT,
        image TEXT, is_active INTEGER DEFAULT 0,
        part_name TEXT, part_responsible TEXT,
        FOREIGN KEY (project_id) REFERENCES project(id) ON DELETE CASCADE
    )
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS project_parts(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        project_id INTEGER, part_name TEXT, created_at TEXT,
        FOREIGN KEY (project_id) REFERENCES project(id) ON DELETE CASCADE
    )
    """)
    # ذخیره تاریخ نصب برای لایسنس
    cur.execute("SELECT value FROM settings WHERE key='install_date'")
    row = cur.fetchone()
    if not row:
        cur.execute("INSERT INTO settings (key, value) VALUES ('install_date', ?)", (datetime.now().strftime("%Y-%m-%d"),))
        conn.commit()
    conn.close()

init_db()

class LicenseManager:
    @staticmethod
    def generate_user_code():
        # فایل پایدار در حافظه عمومی (حتی بعد از حذف نصب باقی می‌ماند)
        code_file = '/sdcard/.mhb_user_code'
        if os.path.exists(code_file):
            try:
                with open(code_file, 'r') as f:
                    code = f.read().strip()
                    if code and code.isdigit() and len(code) == 4:
                        return code
            except:
                pass
        
        # اگر فایل نبود، یک کد جدید بساز و ذخیره کن
        try:
            node = uuid.getnode()
            new_code = str(node)[-4:].zfill(4)
        except:
            new_code = "0000"
        
        # ذخیره در فایل (اگر امکان‌پذیر باشد)
        try:
            with open(code_file, 'w') as f:
                f.write(new_code)
        except:
            pass
        
        # همچنین در دیتابیس ذخیره کن (برای سازگاری)
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT value FROM settings WHERE key='current_user_code'")
        row = cur.fetchone()
        if row:
            cur.execute("UPDATE settings SET value=? WHERE key='current_user_code'", (new_code,))
        else:
            cur.execute("INSERT INTO settings (key, value) VALUES ('current_user_code', ?)", (new_code,))
        conn.commit()
        conn.close()
        return new_code

    @staticmethod
    def generate_verification_code(phone, user_code):
        try:
            phone_str = str(phone).strip()
            phone_last4 = phone_str[-4:] if len(phone_str) >= 4 else phone_str.zfill(4)
            phone_num = int(phone_last4) if phone_last4.isdigit() else 0
            user_code_num = int(user_code) if user_code.isdigit() else 0
            total = user_code_num + (phone_num * 2)
            return str(total)[-4:].zfill(4)
        except:
            return "0000"

    @staticmethod
    def generate_renewal_code(phone, user_code):
        try:
            phone_str = str(phone).strip()
            phone_num = int(phone_str) if phone_str.isdigit() else 0
            user_code_num = int(user_code) if user_code.isdigit() else 0
            total = (phone_num * 3) + user_code_num
            return str(total)[-4:].zfill(4)
        except:
            return "0000"

# ============================================================
# توابع بررسی لایسنس
# ============================================================
def get_license_status():
    conn = get_db()
    cur = conn.cursor()
    cur.execute("SELECT value FROM settings WHERE key='install_date'")
    row = cur.fetchone()
    conn.close()
    if not row:
        install_date = datetime.now()
    else:
        try:
            install_date = datetime.strptime(row['value'], "%Y-%m-%d")
        except:
            install_date = datetime.now()
    current_date = datetime.now()
    days_used = (current_date - install_date).days
    expired = days_used > LICENSE_EXPIRY_DAYS
    return expired, days_used

# ============================================================
# HTML
# ============================================================
HTML_TEMPLATE = """
<!DOCTYPE html>
<html dir="rtl" lang="fa">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>MHB Inspection</title>
    <style>
        body { font-family: Tahoma, sans-serif; background: #f0f2f5; margin: 0; padding: 10px; direction: rtl; }
        .main-container { max-width: 950px; margin: auto; background: white; border-radius: 12px; padding: 15px; box-shadow: 0 4px 10px rgba(0,0,0,0.1); }
        h1 { color: #2c3e50; text-align: center; font-size: 20px; margin-bottom: 15px; }
        .tabs { display: flex; flex-wrap: wrap; border-bottom: 2px solid #ddd; margin-top: 10px; }
        .tab { padding: 10px 15px; cursor: pointer; background: #eee; margin-left: 2px; margin-bottom: 2px; border-radius: 5px 5px 0 0; font-weight: bold; color: #555; flex-grow: 1; text-align: center; font-size: 14px; }
        .tab.active { background: #3498db; color: white; }
        .tab.disabled { background: #ccc; color: #888; cursor: not-allowed; }
        .tab-content { display: none; padding: 15px 10px; }
        .tab-content.active { display: block; }
        .form-group { margin-bottom: 12px; }
        label { display: block; font-weight: bold; margin-bottom: 5px; text-align: right; font-size: 14px; }
        input[type="text"], input[type="number"], input[type="date"], textarea, select {
            width: 100%; padding: 8px; border: 1px solid #ccc; border-radius: 4px;
            box-sizing: border-box; text-align: right; direction: rtl; font-size: 14px;
        }
        .btn-sky { background: #87CEEB; color: #000; padding: 8px 15px; border: none; border-radius: 4px; cursor: pointer; font-weight: bold; font-size: 14px; width: 100%; margin-top: 5px; transition: 0.3s; }
        .btn-sky:hover { background: #5B9BD5; color: white; }
        .btn-orange { background: #FF8C00; color: black; padding: 8px 15px; border: none; border-radius: 4px; cursor: pointer; font-weight: bold; font-size: 14px; width: 100%; margin-top: 5px; }
        .btn-red { background: #e74c3c; color: white; }
        .btn-small { padding: 2px 8px; font-size: 10px; width: auto; margin: 0; }
        .btn-green { background: #2ecc71; color: white; padding: 8px 15px; border: none; border-radius: 4px; cursor: pointer; font-weight: bold; font-size: 14px; width: 100%; margin-top: 5px; transition: 0.3s; }
        .btn-green:hover { background: #27ae60; }

        .row-images { display: flex; flex-wrap: wrap; gap: 8px; margin-bottom: 10px; }
        .img-field { flex: 1; min-width: 100px; border: 1px solid #ddd; padding: 8px; border-radius: 4px; }
        .img-preview { max-width: 100%; height: auto; max-height: 80px; display: block; margin: 5px auto; border-radius: 4px; }
        input[type="file"] { font-size: 11px; padding: 3px; width: 100%; box-sizing: border-box; }

        .collapsible-header { background: #e9ecef; padding: 12px 15px; cursor: pointer; font-weight: bold; border: 1px solid #dee2e6; border-radius: 5px; display: flex; justify-content: space-between; align-items: center; margin-bottom: 5px; }
        .collapsible-header:hover { background: #dee2e6; }
        .collapsible-body { padding: 15px; border: 1px solid #dee2e6; border-radius: 5px; display: none; margin-bottom: 15px; }
        .collapsible-body.active { display: block; }

        .project-container { margin-top: 20px; }
        .project-item { background: #f9f9f9; margin-bottom: 10px; border: 1px solid #eee; border-radius: 4px; overflow: hidden; }
        .project-header { background: #e9ecef; padding: 12px 15px; cursor: pointer; font-weight: bold; border-bottom: 1px solid #dee2e6; display: flex; justify-content: space-between; align-items: center; }
        .project-header:hover { background: #dee2e6; }
        .project-body { padding: 15px; display: none; border-top: 1px solid #eee; }
        .project-body.active { display: block; }

        .status-msg { text-align: center; padding: 8px; margin: 10px 0; border-radius: 4px; font-size: 14px; }
        .success { background: #d4edda; color: #155724; border: 1px solid #c3e6cb; }
        .error { background: #f8d7da; color: #721c24; border: 1px solid #f5c6cb; }
        .info { background: #d1ecf1; color: #0c5460; border: 1px solid #bee5eb; }
        .warning { background: #fff3cd; color: #856404; border: 1px solid #ffeeba; }
        .flex { display: flex; gap: 10px; align-items: center; flex-wrap: wrap; }
        .flex-right { justify-content: flex-end; }

        #loading-modal {
            display: none;
            position: fixed; top: 0; left: 0; width: 100%; height: 100%;
            background: rgba(0,0,0,0.7); z-index: 9999;
            justify-content: center; align-items: center; flex-direction: column;
        }
        .loader {
            width: 80px; height: 80px;
            border: 8px solid #f3f3f3; border-top: 8px solid #3498db;
            border-radius: 50%; animation: spin 1s linear infinite;
        }
        @keyframes spin { 0% { transform: rotate(0deg); } 100% { transform: rotate(360deg); } }
        .loader-text { color: white; margin-top: 20px; font-size: 18px; font-weight: bold; }

        .cc-container { display: flex; flex-direction: column; gap: 5px; margin-top: 10px; }
        .cc-row { display: flex; gap: 10px; align-items: center; }
        .cc-row input { flex: 1; padding: 5px; font-size: 14px; border: 1px solid #ccc; border-radius: 4px; text-align: right; direction: rtl; }
        .cc-row button { padding: 5px 10px; background: #27ae60; color: white; border: none; border-radius: 4px; cursor: pointer; }

        .edit-images-container { display: flex; flex-wrap: wrap; gap: 15px; margin-top: 10px; }
        .edit-image-item { border: 1px solid #ddd; padding: 5px; border-radius: 4px; text-align: center; background: #f9f9f9; }
        .edit-image-item img { max-width: 100px; max-height: 80px; display: block; margin: 0 auto; }
        .edit-image-item .delete-btn { background: none; border: none; cursor: pointer; color: #e74c3c; font-size: 18px; padding: 2px 8px; }
        .edit-image-item .delete-btn:hover { color: #c0392b; }

        .project-actions { display: flex; gap: 5px; align-items: center; }
        .project-edit-form { margin-top: 10px; padding: 10px; border: 1px solid #ddd; border-radius: 4px; background: #fefefe; display: none; }
        .project-edit-form.active { display: block; }

        @media (max-width: 600px) {
            .tab { font-size: 12px; padding: 6px 10px; flex-grow: 0; width: 45%; }
            h1 { font-size: 18px; }
            .btn-sky { font-size: 13px; }
        }
    </style>
</head>
<body>
    {% if license_expired %}
        <!-- صفحه قفل لایسنس -->
        <div class="main-container" style="text-align:center; padding:40px;">
            <div style="font-size:60px;">🔒</div>
            <h2>اشتراک شما به پایان رسیده است</h2>
            <p>لطفاً برای تمدید اشتراک، با پشتیبانی تماس بگیرید و رمز تمدید را وارد کنید.</p>
            <form action="/renew_license" method="POST" style="max-width:300px; margin:20px auto;">
                <div class="form-group">
                    <label>رمز تمدید</label>
                    <input type="text" name="renewal_code" maxlength="4" style="text-align:center; font-size:18px;" required>
                </div>
                <button type="submit" class="btn-sky">تمدید اشتراک</button>
            </form>
            {% with messages = get_flashed_messages(with_categories=true) %}
                {% if messages %}{% for category, message in messages %}<div class="status-msg {{ category }}">{{ message }}</div>{% endfor %}{% endif %}
            {% endwith %}
        </div>
    {% else %}
    <div class="main-container">
        <h1>نرم افزار جامع مدیریت هوشمند بازدید MHB</h1>
        <div class="tabs">
            <div class="tab {% if not is_verified %}disabled{% endif %}" onclick="switchTab('tab-user')">کاربر</div>
            <div class="tab {% if not is_verified %}disabled{% endif %}" onclick="switchTab('tab-project')">پروژه</div>
            <div class="tab {% if not is_verified %}disabled{% endif %}" onclick="switchTab('tab-defect')">انطباق</div>
            <div class="tab {% if not is_verified or not can_access_report %}disabled{% endif %}" onclick="switchTab('tab-report')">گزارش</div>
        </div>

        <!-- تب کاربر -->
        <div id="tab-user" class="tab-content active">
            <h2>مدیریت اطلاعات کاربر</h2>
            
            <div class="collapsible-header" onclick="toggleCollapse(this)">
                <span>👤 اطلاعات کاربر</span>
                <span style="font-size:12px; color:#666;">🖱️ کلیک کنید</span>
            </div>
            <div class="collapsible-body">
                <form action="/save_user" method="POST" enctype="multipart/form-data">
                    <div class="form-group"><label>نام و نام خانوادگی</label><input type="text" name="fullname" value="{{ user.fullname }}"></div>
                    <div class="form-group"><label>سمت</label><input type="text" name="position" value="{{ user.position }}"></div>
                    <div class="form-group"><label>کد شناسایی</label><input type="text" name="license" value="{{ user.license }}"></div>
                    <div class="form-group"><label>شماره تماس</label><input type="text" name="phone" value="{{ user.phone }}"></div>
                    <div class="row-images">
                        <div class="img-field"><label>لوگو</label><div style="display:flex; gap:5px;"><input type="file" name="logo" accept="image/*" capture="environment"><button type="submit" formaction="/remove_user_logo" class="btn-sky btn-small" style="background:#FF8C00;">🗑️</button></div>{% if user.logo_path %}<img src="{{ url_for('static', filename='uploads/' + user.logo_path.split('/')[-1]) }}" class="img-preview">{% endif %}</div>
                        <div class="img-field"><label>امضا</label><div style="display:flex; gap:5px;"><input type="file" name="signature" accept="image/*" capture="environment"><button type="submit" formaction="/remove_user_sig" class="btn-sky btn-small" style="background:#FF8C00;">🗑️</button></div>{% if user.signature_path %}<img src="{{ url_for('static', filename='uploads/' + user.signature_path.split('/')[-1]) }}" class="img-preview">{% endif %}</div>
                        <div class="img-field"><label>مهر</label><div style="display:flex; gap:5px;"><input type="file" name="stamp" accept="image/*" capture="environment"><button type="submit" formaction="/remove_user_stamp" class="btn-sky btn-small" style="background:#FF8C00;">🗑️</button></div>{% if user.stamp_path %}<img src="{{ url_for('static', filename='uploads/' + user.stamp_path.split('/')[-1]) }}" class="img-preview">{% endif %}</div>
                    </div>
                    <button type="submit" class="btn-sky">💾 ذخیره اطلاعات</button>
                </form>
            </div>

            <!-- فعال‌سازی: کد کاربری نمایش داده می‌شود، فرمول مخفی است -->
            <div class="collapsible-header" onclick="toggleCollapse(this)">
                <span>🔑 فعال‌سازی</span>
                <span style="font-size:12px; color:#666;">🖱️ کلیک کنید</span>
            </div>
            <div class="collapsible-body">
                <form action="/verify_code" method="POST">
                    <div class="activation-row" style="display:flex; flex-wrap:wrap; gap:10px;">
                        <div class="form-group" style="flex:1;"><label>کد کاربری دستگاه</label><input type="text" value="{{ user_code }}" readonly style="background:#e9ecef; text-align:center;"></div>
                        <div class="form-group" style="flex:1;"><label>رمز تایید</label><input type="password" name="entered_code" placeholder="****" maxlength="4" style="text-align:center;"></div>
                        <div style="flex: 0 0 100px; display:flex; align-items:center;"><button type="submit" class="btn-sky" style="width:100%;">تأیید</button></div>
                    </div>
                </form>
                {% if is_verified %}<div class="status-msg success" style="margin-top:8px;">✅ دسترسی فعال است!</div>{% endif %}
                
                <!-- نمایش وضعیت اشتراک -->
                <div class="status-msg info" style="margin-top:8px;">
                    📅 وضعیت اشتراک: <b>{{ license_status }}</b> | روزهای باقی‌مانده: <b>{{ license_days_left }}</b> روز
                </div>
            </div>
        </div>

        <!-- تب پروژه -->
        <div id="tab-project" class="tab-content">
            <h2>مدیریت پروژه‌ها</h2>
            
            <div class="collapsible-header" onclick="toggleCollapse(this)">
                <span>➕ ثبت پروژه جدید</span>
                <span style="font-size:12px; color:#666;">🖱️ کلیک کنید</span>
            </div>
            <div class="collapsible-body">
                <form action="/add_project" method="POST" enctype="multipart/form-data">
                    <div class="form-group"><label>شناسه پروژه (یکتا)</label><input type="text" name="report_no"></div>
                    <div class="form-group"><label>نام پروژه</label><input type="text" name="project_name"></div>
                    <div class="form-group"><label>مدیر</label><input type="text" name="employer"></div>
                    <div class="form-group"><label>آدرس</label><input type="text" name="address"></div>
                    <div class="form-group" style="display:flex; gap:5px; align-items:center;">
                        <label style="flex:0 0 80px;">لوگو:</label>
                        <input type="file" name="project_logo" accept="image/*" style="flex-grow:1;">
                        <button type="submit" formaction="/remove_project_logo" class="btn-orange btn-small">🗑️</button>
                    </div>
                    <button type="submit" class="btn-sky">ذخیره پروژه</button>
                </form>
            </div>

            <h3 style="margin-top:20px;">پروژه‌های ثبت شده</h3>
            <div class="project-container">
            {% for p in projects %}
                <div class="project-item">
                    <div class="project-header" onclick="toggleProject(this)">
                        <span><b>{{ p.project_name }}</b> ({{ p.report_no }})</span>
                        <div class="project-actions" onclick="event.stopPropagation();">
                            <button class="btn-sky btn-small" style="background:#f39c12;" onclick="toggleProjectEdit({{ p.id }})">✏️</button>
                            <form action="/delete_project/{{ p.id }}" method="POST" style="display:inline;" onsubmit="return confirm('آیا از حذف پروژه {{ p.project_name }} و تمام اطلاعات مربوط به آن اطمینان دارید؟');">
                                <button type="submit" class="btn-red btn-small">🗑️</button>
                            </form>
                        </div>
                    </div>
                    <div class="project-body">
                        
                        <div class="project-edit-form" id="edit-project-{{ p.id }}">
                            <form action="/edit_project/{{ p.id }}" method="POST" enctype="multipart/form-data">
                                <div class="form-group"><label>شناسه پروژه (یکتا)</label><input type="text" name="report_no" value="{{ p.report_no }}" required></div>
                                <div class="form-group"><label>نام پروژه</label><input type="text" name="project_name" value="{{ p.project_name }}" required></div>
                                <div class="form-group"><label>مدیر</label><input type="text" name="employer" value="{{ p.employer }}"></div>
                                <div class="form-group"><label>آدرس</label><input type="text" name="address" value="{{ p.address }}"></div>
                                <div class="form-group" style="display:flex; gap:5px; align-items:center;">
                                    <label style="flex:0 0 80px;">لوگو جدید:</label>
                                    <input type="file" name="project_logo" accept="image/*" style="flex-grow:1;">
                                </div>
                                <button type="submit" class="btn-sky" style="background:#2ecc71;">💾 ذخیره تغییرات</button>
                            </form>
                        </div>

                        <div class="collapsible-header" onclick="toggleCollapse(this)">
                            <span>⚙️ مدیریت قسمت‌ها/مراحل این پروژه</span>
                            <span style="font-size:12px; color:#666;">🖱️ کلیک کنید</span>
                        </div>
                        <div class="collapsible-body">
                            <div style="display:flex; gap:10px; margin-bottom:10px;">
                                <form action="/add_part" method="POST" style="display:flex; flex-grow:1; gap:5px;">
                                    <input type="hidden" name="project_id" value="{{ p.id }}">
                                    <input type="text" name="part_name" placeholder="نام قسمت جدید" style="flex-grow:1;">
                                    <button type="submit" class="btn-sky btn-small">➕</button>
                                </form>
                            </div>
                            <div style="margin-bottom:10px;">
                                {% for part in project_parts[p.id] %}
                                <div style="display:flex; justify-content:space-between; align-items:center; padding:4px 0; border-bottom:1px solid #f0f0f0;">
                                    <span id="part-name-{{ part.id }}">{{ part.part_name }}</span>
                                    <div>
                                        <form action="/edit_part" method="POST" style="display:inline;" id="edit-form-{{ part.id }}">
                                            <input type="hidden" name="part_id" value="{{ part.id }}">
                                            <input type="text" name="new_name" value="{{ part.part_name }}" style="display:none; width:100px; font-size:12px; padding:2px;" id="part-input-{{ part.id }}">
                                            <button type="button" class="btn-sky btn-small" style="background:#f39c12;" onclick="toggleEdit({{ part.id }})" id="edit-btn-{{ part.id }}">✏️</button>
                                            <button type="submit" class="btn-sky btn-small" style="display:none; background:#2ecc71;" id="save-btn-{{ part.id }}">💾</button>
                                        </form>
                                        <form action="/delete_part" method="POST" style="display:inline;">
                                            <input type="hidden" name="part_id" value="{{ part.id }}">
                                            <button type="submit" class="btn-red btn-small">🗑️</button>
                                        </form>
                                    </div>
                                </div>
                                {% endfor %}
                            </div>
                        </div>

                        <div class="collapsible-header" onclick="toggleCollapse(this)">
                            <span>📋 گزارش‌های قبلی این پروژه</span>
                            <span style="font-size:12px; color:#666;">🖱️ کلیک کنید</span>
                        </div>
                        <div class="collapsible-body">
                            {% set ns = namespace(found=false) %}
                            {% for d in defects if d.project_id == p.id %}
                                {% set ns.found = true %}
                                <div style="display:flex; justify-content:space-between; align-items:center; padding:4px 0; border-bottom:1px solid #f0f0f0;">
                                    <div style="flex:1; text-align:right;">
                                        <b>{{ d.report_no }}</b> - {{ d.part_name }}<br>
                                        <span style="font-size:12px; color:#666;">{{ d.title or 'بدون عنوان' }}</span>
                                    </div>
                                    <div style="display:flex; gap:5px;">
                                        <form action="/edit_defect/{{ d.id }}" method="POST" style="display:inline;">
                                            <button type="submit" class="btn-sky btn-small" style="background:#f39c12;">✏️</button>
                                        </form>
                                        <form action="/delete_defect/{{ d.id }}" method="POST" style="display:inline;">
                                            <button type="submit" class="btn-red btn-small" onclick="return confirm('آیا از حذف این گزارش اطمینان دارید؟');">🗑️</button>
                                        </form>
                                        <form action="/set_report/{{ d.id }}" method="POST" style="display:inline;">
                                            <button type="submit" class="btn-sky btn-small" style="background:#2ecc71;">📄</button>
                                        </form>
                                    </div>
                                </div>
                            {% endfor %}
                            {% if not ns.found %}
                                <p style="color:#999; font-size:13px;">هیچ گزارشی برای این پروژه ثبت نشده است.</p>
                            {% endif %}
                        </div>

                        <div class="collapsible-header" onclick="toggleCollapse(this)">
                            <span>📝 ثبت بازدید جدید</span>
                            <span style="font-size:12px; color:#666;">🖱️ کلیک کنید</span>
                        </div>
                        <div class="collapsible-body">
                            <form action="/add_visit" method="POST">
                                <input type="hidden" name="project_id" value="{{ p.id }}">
                                <div class="form-group">
                                    <label>انتخاب قسمت/مرحله</label>
                                    <select name="part_name" required>
                                        <option value="">انتخاب کنید</option>
                                        {% for part in project_parts[p.id] %}
                                        <option value="{{ part.part_name }}">{{ part.part_name }}</option>
                                        {% endfor %}
                                    </select>
                                </div>
                                <div class="form-group"><label>شماره گزارش (یکتا)</label><input type="text" name="report_no_visit" required></div>
                                <div class="form-group"><label>تاریخ</label><input type="date" name="visit_date"></div>
                                <div class="form-group"><label>نوبت بازدید</label><input type="text" name="stage" placeholder="اول، دوره‌ای..."></div>
                                <button type="submit" class="btn-sky">ثبت بازدید</button>
                            </form>
                        </div>

                    </div>
                </div>
            {% endfor %}
            </div>
        </div>

        <!-- تب انطباق -->
        <div id="tab-defect" class="tab-content">
            {% if current_visit and not edit_defect %}
                <div class="status-msg success" style="font-size:16px;">
                    <b>📌 شما در حال ثبت مورد شماره {{ defect_counter }} از پروژه {{ current_visit.project_name }} با شماره گزارش {{ current_visit.report_no }} هستید.</b>
                </div>
                <h2>ثبت انطباق یا عدم انطباق (چند تصویر)</h2>
                <form action="/add_defect" method="POST" enctype="multipart/form-data">
                    <input type="hidden" name="defect_id" value="{{ current_visit.defect_id }}">
                    <div class="form-group"><label>عنوان</label><input type="text" name="title"></div>
                    <div class="form-group"><label>استاندارد</label><input type="text" name="standard"></div>
                    <div class="form-group"><label>شرح</label><textarea name="description" rows="3"></textarea></div>
                    <div class="form-group"><label>تصاویر (چند فایل را با هم انتخاب کنید)</label><input type="file" name="image" multiple accept="image/*" capture="environment"></div>
                    <div style="display:flex; gap:10px;">
                        <button type="submit" name="action" value="continue" class="btn-sky btn-small" style="width:auto;">➕ ثبت جدید</button>
                        <button type="submit" name="action" value="finalize" class="btn-green" style="width:auto;">✅ ثبت و تولید</button>
                    </div>
                </form>
            {% elif edit_defect %}
                <div class="status-msg info" style="font-size:16px;">
                    <b>✏️ شما در حال ویرایش گزارش {{ edit_defect.report_no }} از پروژه {{ edit_defect.project_name }} هستید.</b>
                </div>
                <h2>ویرایش انطباق یا عدم انطباق (چند تصویر)</h2>
                <form action="/update_defect/{{ edit_defect.id }}" method="POST" enctype="multipart/form-data" id="edit-defect-form">
                    <input type="hidden" name="defect_id" value="{{ edit_defect.id }}">
                    <div class="form-group"><label>عنوان</label><input type="text" name="title" value="{{ edit_defect.title or '' }}"></div>
                    <div class="form-group"><label>استاندارد</label><input type="text" name="standard" value="{{ edit_defect.standard or '' }}"></div>
                    <div class="form-group"><label>شرح</label><textarea name="description" rows="3">{{ edit_defect.description or '' }}</textarea></div>
                    
                    <div class="form-group">
                        <label>تصاویر موجود</label>
                        {% set titles = edit_defect.title.split(' ||| ') if edit_defect.title else [] %}
                        {% set images_groups = edit_defect.image.split(' ||| ') if edit_defect.image else [] %}
                        {% set max_groups = titles|length if titles|length > images_groups|length else images_groups|length %}
                        {% if max_groups == 0 %}
                            <p style="color:#999;">هیچ موردی ثبت نشده است.</p>
                        {% else %}
                            {% for i in range(max_groups) %}
                                {% set group_title = titles[i] if i < titles|length else '' %}
                                {% set group_images = [] %}
                                {% if i < images_groups|length %}
                                    {% for img in images_groups[i].split('&&&') %}
                                        {% if img.strip() %}
                                            {% set _ = group_images.append(img.strip()) %}
                                        {% endif %}
                                    {% endfor %}
                                {% endif %}
                                <div style="border:1px solid #ddd; padding:10px; margin-bottom:10px; border-radius:4px;">
                                    <h4>مورد شماره {{ i+1 }}: {{ group_title if group_title else 'بدون عنوان' }}</h4>
                                    <div class="edit-images-container">
                                        {% if group_images|length > 0 %}
                                            {% for img_path in group_images %}
                                                <div class="edit-image-item">
                                                    <img src="{{ url_for('static', filename='uploads/' + img_path.split('/')[-1]) }}" alt="تصویر">
                                                    <button type="button" class="delete-btn" onclick="deleteImage({{ i }}, {{ loop.index0 }})">🗑️</button>
                                                    <input type="hidden" name="delete_images[]" id="del-{{ i }}-{{ loop.index0 }}" value="">
                                                </div>
                                            {% endfor %}
                                        {% else %}
                                            <p style="color:#999; width:100%;">هیچ تصویری برای این مورد ثبت نشده است.</p>
                                        {% endif %}
                                    </div>
                                    <div style="margin-top:5px;">
                                        <label>افزودن تصویر جدید به این مورد:</label>
                                        <input type="file" name="new_images_group_{{ i }}" multiple accept="image/*" capture="environment">
                                    </div>
                                </div>
                            {% endfor %}
                        {% endif %}
                    </div>

                    <button type="submit" class="btn-green" style="width:100%;">💾 ذخیره و تولید گزارش</button>
                </form>
            {% else %}
                <div class="status-msg error" style="font-size:16px;">
                    ⛔ لطفا در تب پروژه "ثبت بازدید جدید" انجام دهید.
                </div>
                <p style="text-align:center; margin-top:20px;">
                    <a href="{{ url_for('index', active_tab='tab-project') }}" class="btn-sky" style="display:inline-block; width:auto; padding:8px 25px;">بازگشت به پروژه‌ها</a>
                </p>
            {% endif %}
        </div>

        <!-- تب گزارش نهایی -->
        <div id="tab-report" class="tab-content">
            {% if report_just_generated %}
                <!-- پیام بعد از تولید گزارش -->
                <div class="status-msg info" style="font-size:16px; margin-bottom:20px;">
                    <b>📌 گزارش با موفقیت تولید شد.</b><br>
                    لطفا جهت تولید گزارش بعدی از تب پروژه اقدام کنید.
                    <br><br>
                    <a href="{{ url_for('index', active_tab='tab-project') }}" class="btn-sky" style="display:inline-block; width:auto; padding:8px 25px;">بازگشت به پروژه‌ها</a>
                </div>
            {% elif current_report %}
                <div class="status-msg success" style="font-size:16px;">
                    <b>📌 شما در حال تولید گزارش نهایی با این مشخصات هستید:</b><br>
                    شماره گزارش: {{ current_report.report_no }}<br>
                    پروژه: {{ current_report.project_name }} | قسمت: {{ current_report.part_name }} | نوبت: {{ current_report.stage }}
                </div>

                <h2>تولید گزارش نهایی</h2>
                <form id="reportForm" action="/generate_report" method="POST" onsubmit="showLoadingModal(); return true;">
                    {% if current_report and current_report.defect_id %}
                        <input type="hidden" name="report_id" value="{{ current_report.defect_id }}">
                    {% endif %}
                    <div class="form-group"><label>سرتیتر گزارش</label><textarea name="header" rows="2"></textarea></div>
                    <div class="form-group"><label>نتیجه بازدید</label><textarea name="result" rows="2"></textarea></div>

                    <div class="form-group">
                        <label>رونوشت (کپی برای ارسال)</label>
                        <div class="cc-container" id="cc-container">
                            <div class="cc-row">
                                <input type="text" name="cc[]" placeholder="نام شخص یا سازمان...">
                                <button type="button" onclick="addCCField()" class="btn-sky btn-small" style="background:#27ae60;">+</button>
                            </div>
                        </div>
                    </div>

                    <div class="flex flex-right"><label><input type="checkbox" name="pdf" checked> PDF</label><label><input type="checkbox" name="word"> Word</label></div>
                    <button type="submit" class="btn-sky">📄 تولید گزارش</button>
                </form>
            {% else %}
                <!-- عدم دسترسی مستقیم -->
                <div class="status-msg error" style="font-size:16px;">
                    ⛔ لطفا جهت تولید گزارش از تب پروژه اقدام کنید.
                </div>
                <p style="text-align:center; margin-top:20px;">
                    <a href="{{ url_for('index', active_tab='tab-project') }}" class="btn-sky" style="display:inline-block; width:auto; padding:8px 25px;">بازگشت به پروژه‌ها</a>
                </p>
            {% endif %}
        </div>

        {% with messages = get_flashed_messages(with_categories=true) %}
          {% if messages %}{% for category, message in messages %}<div class="status-msg {{ category }}">{{ message | safe }}</div>{% endfor %}{% endif %}
        {% endwith %}
    </div>
    {% endif %}

    <script>
        var can_access_report = {{ 'true' if can_access_report else 'false' }};
        
        function switchTab(tabId) {
            if (tabId === 'tab-report' && !can_access_report) {
                window.location.href = '?active_tab=tab-project';
                return;
            }
            const tab = document.querySelector(`.tab[onclick*="switchTab('${tabId}')"]`);
            if (tab && tab.classList.contains('disabled')) {
                return;
            }
            document.querySelectorAll('.tab-content').forEach(el => el.classList.remove('active'));
            document.querySelectorAll('.tab').forEach(el => el.classList.remove('active'));
            document.getElementById(tabId).classList.add('active');
            if(tab) tab.classList.add('active');
        }
        
        function toggleCollapse(header) {
            var body = header.nextElementSibling;
            if (body.classList.contains('active')) body.classList.remove('active');
            else body.classList.add('active');
        }
        function toggleProject(header) {
            var body = header.nextElementSibling;
            if (body.classList.contains('active')) body.classList.remove('active');
            else body.classList.add('active');
        }

        function toggleEdit(partId) {
            var span = document.getElementById('part-name-'+partId);
            var input = document.getElementById('part-input-'+partId);
            var editBtn = document.getElementById('edit-btn-'+partId);
            var saveBtn = document.getElementById('save-btn-'+partId);
            if (span.style.display !== 'none') {
                span.style.display = 'none';
                input.style.display = 'inline-block';
                editBtn.style.display = 'none';
                saveBtn.style.display = 'inline-block';
                input.value = span.textContent.trim();
            } else {
                span.style.display = 'inline';
                input.style.display = 'none';
                editBtn.style.display = 'inline-block';
                saveBtn.style.display = 'none';
            }
        }

        function toggleProjectEdit(projectId) {
            var form = document.getElementById('edit-project-'+projectId);
            if (form.classList.contains('active')) {
                form.classList.remove('active');
            } else {
                form.classList.add('active');
            }
        }

        function deleteImage(groupIndex, imageIndex) {
            var hiddenInput = document.getElementById('del-' + groupIndex + '-' + imageIndex);
            if (hiddenInput) {
                hiddenInput.value = groupIndex + ':' + imageIndex;
                document.getElementById('edit-defect-form').submit();
            }
        }

        document.addEventListener('DOMContentLoaded', function() {
            const urlParams = new URLSearchParams(window.location.search);
            const activeTab = urlParams.get('active_tab');
            if (activeTab) switchTab(activeTab);
        });

        function showLoadingModal() {
            const modal = document.getElementById('loading-modal');
            const percentText = document.getElementById('percent-text');
            if (modal) modal.style.display = 'flex';
            if (percentText) {
                let p = 0;
                const interval = setInterval(() => {
                    p += 2;
                    if(p > 99) p = 99;
                    percentText.innerText = `در حال تولید گزارش... ${p}%`;
                }, 300);
            }
            return true;
        }

        function addCCField() {
            const container = document.getElementById('cc-container');
            const newRow = document.createElement('div');
            newRow.className = 'cc-row';
            newRow.innerHTML = `
                <input type="text" name="cc[]" placeholder="نام شخص یا سازمان...">
                <button type="button" onclick="removeCCField(this)" class="btn-red btn-small" style="background:#e74c3c;">✖</button>
            `;
            container.appendChild(newRow);
        }

        function removeCCField(btn) {
            btn.parentElement.remove();
        }
    </script>
</body>
</html>
"""

# ============================================================
# مسیرهای Flask
# ============================================================

@app.route('/')
def index():
    active_tab = request.args.get('active_tab', 'tab-user')

    # بررسی وضعیت لایسنس
    expired, days_used = get_license_status()
    if expired:
        return render_template_string(HTML_TEMPLATE, license_expired=True)

    if active_tab == 'tab-report':
        if not session.get('current_report') and not session.get('report_just_generated'):
            flash('لطفا جهت تولید گزارش از تب پروژه اقدام کنید.', 'info')
            return redirect(url_for('index', active_tab='tab-project'))

    if active_tab == 'tab-defect':
        current_visit = session.get('current_visit')
        edit_defect_id = session.get('edit_defect_id')
        
        if not current_visit and not edit_defect_id:
            flash('لطفا در تب پروژه "ثبت بازدید جدید" انجام دهید.', 'info')
            return redirect(url_for('index', active_tab='tab-project'))

        if current_visit:
            conn = get_db()
            cur = conn.cursor()
            cur.execute("SELECT id FROM defect WHERE id=?", (current_visit.get('defect_id'),))
            exists = cur.fetchone()
            conn.close()
            if not exists:
                session.pop('current_visit', None)
                flash('بازدید مورد نظر منقضی شده است. لطفاً دوباره ثبت بازدید کنید.', 'info')
                return redirect(url_for('index', active_tab='tab-project'))

        if edit_defect_id:
            conn = get_db()
            cur = conn.cursor()
            cur.execute("SELECT id FROM defect WHERE id=?", (edit_defect_id,))
            exists = cur.fetchone()
            conn.close()
            if not exists:
                session.pop('edit_defect_id', None)
                flash('گزارش مورد نظر یافت نشد.', 'error')
                return redirect(url_for('index', active_tab='tab-project'))

    conn = get_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM expert LIMIT 1")
    user = cur.fetchone()
    cur.execute("SELECT * FROM project ORDER BY id DESC")
    projects = cur.fetchall()
    cur.execute("SELECT d.*, p.project_name FROM defect d JOIN project p ON d.project_id = p.id ORDER BY d.id DESC")
    defects = cur.fetchall()
    
    project_parts = {}
    for p in projects:
        cur.execute("SELECT * FROM project_parts WHERE project_id=? ORDER BY part_name", (p['id'],))
        project_parts[p['id']] = cur.fetchall()

    cur.execute("SELECT value FROM settings WHERE key='is_verified'")
    res = cur.fetchone()
    is_verified = res and res['value'] == '1'

    edit_defect = None
    edit_defect_id = session.get('edit_defect_id')
    if edit_defect_id:
        cur.execute("""
            SELECT d.*, p.project_name 
            FROM defect d 
            JOIN project p ON d.project_id = p.id 
            WHERE d.id=?
        """, (edit_defect_id,))
        edit_defect = cur.fetchone()
        if not edit_defect:
            session.pop('edit_defect_id', None)

    conn.close()

    current_visit = session.get('current_visit')
    current_report = session.get('current_report')
    defect_counter = session.get('defect_counter', 1)
    report_just_generated = session.get('report_just_generated', False)
    
    can_access_report = bool(current_report or report_just_generated)

    license_days_left = max(0, LICENSE_EXPIRY_DAYS - days_used)
    license_status = "فعال" if not expired else "منقضی"

    user_code = LicenseManager.generate_user_code()

    return render_template_string(HTML_TEMPLATE, user=dict(user) if user else {}, projects=projects, defects=defects, project_parts=project_parts, is_verified=is_verified, current_visit=current_visit, current_report=current_report, edit_defect=edit_defect, defect_counter=defect_counter, report_just_generated=report_just_generated, license_expired=False, can_access_report=can_access_report, user_code=user_code, license_days_left=license_days_left, license_status=license_status)

@app.route('/renew_license', methods=['POST'])
def renew_license():
    renewal_code = request.form.get('renewal_code', '').strip()
    if not renewal_code:
        flash('لطفاً رمز تمدید را وارد کنید.', 'error')
        return redirect(url_for('index'))
    
    user_code = LicenseManager.generate_user_code()
    conn = get_db()
    cur = conn.cursor()
    cur.execute("SELECT phone FROM expert LIMIT 1")
    user = cur.fetchone()
    conn.close()
    
    if not user or not user['phone']:
        flash('لطفاً ابتدا در تب کاربر شماره موبایل خود را ذخیره کنید.', 'error')
        return redirect(url_for('index'))
    
    expected_code = LicenseManager.generate_renewal_code(user['phone'], user_code)
    if renewal_code == expected_code:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("INSERT OR REPLACE INTO settings (key, value) VALUES ('install_date', ?)", (datetime.now().strftime("%Y-%m-%d"),))
        conn.commit()
        conn.close()
        flash('✅ اشتراک با موفقیت تمدید شد!', 'success')
        return redirect(url_for('index'))
    else:
        flash('❌ رمز تمدید اشتباه است!', 'error')
        return redirect(url_for('index'))

@app.route('/dismiss_welcome')
def dismiss_welcome():
    return '', 204

# ============================================================
# کاربر
# ============================================================
@app.route('/save_user', methods=['POST'])
def save_user():
    try:
        fullname, position, license, phone = request.form.get('fullname'), request.form.get('position'), request.form.get('license'), request.form.get('phone')
        def save_uploaded_file(field):
            file = request.files.get(field)
            if file and file.filename:
                filename = secure_filename(f"{field}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.jpg")
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                return filepath
            return None
        logo, sig, stamp = save_uploaded_file('logo'), save_uploaded_file('signature'), save_uploaded_file('stamp')
        conn = get_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM expert")
        cur.execute("INSERT INTO expert (fullname, position, license, phone, logo_path, signature_path, stamp_path) VALUES (?, ?, ?, ?, ?, ?, ?)", (fullname, position, license, phone, logo, sig, stamp))
        conn.commit()
        conn.close()
        flash('اطلاعات کاربر با موفقیت ذخیره شد', 'success')
    except Exception as e:
        flash(f'خطا در ذخیره کاربر: {str(e)}', 'error')
    return redirect(url_for('index'))

@app.route('/remove_user_logo', methods=['POST'])
def remove_user_logo():
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("UPDATE expert SET logo_path = NULL")
        conn.commit()
        conn.close()
        flash('✅ لوگوی کاربر حذف شد.', 'success')
    except Exception as e:
        flash(f'خطا در حذف لوگو: {str(e)}', 'error')
    return redirect(url_for('index'))

@app.route('/remove_user_sig', methods=['POST'])
def remove_user_sig():
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("UPDATE expert SET signature_path = NULL")
        conn.commit()
        conn.close()
        flash('✅ امضای کاربر حذف شد.', 'success')
    except Exception as e:
        flash(f'خطا در حذف امضا: {str(e)}', 'error')
    return redirect(url_for('index'))

@app.route('/remove_user_stamp', methods=['POST'])
def remove_user_stamp():
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("UPDATE expert SET stamp_path = NULL")
        conn.commit()
        conn.close()
        flash('✅ مهر کاربر حذف شد.', 'success')
    except Exception as e:
        flash(f'خطا در حذف مهر: {str(e)}', 'error')
    return redirect(url_for('index'))

@app.route('/verify_code', methods=['POST'])
def verify_code():
    try:
        user_code = LicenseManager.generate_user_code()
        entered_code = request.form.get('entered_code')
        if not entered_code:
            flash('لطفاً رمز تأیید را وارد کنید.', 'error')
            return redirect(url_for('index'))
        
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT phone FROM expert LIMIT 1")
        user = cur.fetchone()
        conn.close()

        if not user or not user['phone'] or len(str(user['phone'])) < 4:
            flash('❌ لطفاً ابتدا در تب "کاربر" شماره موبایل خود را ذخیره کنید.', 'error')
            return redirect(url_for('index'))

        expected_code = LicenseManager.generate_verification_code(user['phone'], user_code)
        
        if entered_code == expected_code:
            conn = get_db()
            cur = conn.cursor()
            cur.execute("INSERT OR REPLACE INTO settings (key, value) VALUES ('is_verified', '1')")
            conn.commit()
            conn.close()
            flash('✅ رمز صحیح است! دسترسی فعال شد.', 'success')
        else:
            flash('❌ رمز وارد شده اشتباه است!', 'error')
    except Exception as e:
        flash(f'خطا: {str(e)}', 'error')
    return redirect(url_for('index'))

# ============================================================
# پروژه (ویرایش و حذف)
# ============================================================
@app.route('/add_project', methods=['POST'])
def add_project():
    try:
        report_no = request.form.get('report_no')
        project_name = request.form.get('project_name')
        employer = request.form.get('employer')
        address = request.form.get('address')
        
        file = request.files.get('project_logo')
        logo_path = ''
        if file and file.filename:
            filename = secure_filename(f"proj_logo_{datetime.now().strftime('%Y%m%d_%H%M%S')}.jpg")
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            logo_path = filepath

        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT id FROM project WHERE report_no=?", (report_no,))
        if cur.fetchone():
            flash('این شناسه پروژه قبلاً ثبت شده است', 'error')
            conn.close()
            return redirect(url_for('index', active_tab='tab-project'))

        cur.execute("INSERT INTO project (report_no, project_name, employer, address, logo_path) VALUES (?, ?, ?, ?, ?)", (report_no, project_name, employer, address, logo_path))
        project_id = cur.lastrowid
        cur.execute("INSERT INTO project_parts (project_id, part_name, created_at) VALUES (?, ?, ?)", (project_id, 'سایر', datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
        conn.commit()
        conn.close()

        project_folder = os.path.join(REPORTS_DIR, re.sub(r'[\\/*?:"<>|]', '_', str(report_no)))
        if not os.path.exists(project_folder):
            try: os.makedirs(project_folder)
            except: pass

        flash(f'✅ پروژه با موفقیت ثبت شد.', 'success')
        return redirect(url_for('index', active_tab='tab-project'))
    except Exception as e:
        flash(f'خطا در ثبت پروژه: {str(e)}', 'error')
        return redirect(url_for('index', active_tab='tab-project'))

@app.route('/edit_project/<int:project_id>', methods=['POST'])
def edit_project(project_id):
    try:
        report_no = request.form.get('report_no')
        project_name = request.form.get('project_name')
        employer = request.form.get('employer')
        address = request.form.get('address')
        
        file = request.files.get('project_logo')
        logo_path = None
        if file and file.filename:
            filename = secure_filename(f"proj_logo_{datetime.now().strftime('%Y%m%d_%H%M%S')}.jpg")
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            logo_path = filepath

        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT id FROM project WHERE report_no=? AND id!=?", (report_no, project_id))
        if cur.fetchone():
            flash('این شناسه پروژه قبلاً ثبت شده است.', 'error')
            conn.close()
            return redirect(url_for('index', active_tab='tab-project'))

        if logo_path:
            cur.execute("UPDATE project SET report_no=?, project_name=?, employer=?, address=?, logo_path=? WHERE id=?", 
                        (report_no, project_name, employer, address, logo_path, project_id))
        else:
            cur.execute("UPDATE project SET report_no=?, project_name=?, employer=?, address=? WHERE id=?", 
                        (report_no, project_name, employer, address, project_id))
        conn.commit()
        conn.close()
        flash('✅ اطلاعات پروژه با موفقیت به‌روزرسانی شد.', 'success')
    except Exception as e:
        flash(f'خطا در ویرایش پروژه: {str(e)}', 'error')
    return redirect(url_for('index', active_tab='tab-project'))

@app.route('/delete_project/<int:project_id>', methods=['POST'])
def delete_project(project_id):
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM project WHERE id=?", (project_id,))
        conn.commit()
        conn.close()
        flash('✅ پروژه و تمام اطلاعات مربوط به آن با موفقیت حذف شد.', 'success')
    except Exception as e:
        flash(f'خطا در حذف پروژه: {str(e)}', 'error')
    return redirect(url_for('index', active_tab='tab-project'))

@app.route('/remove_project_logo', methods=['POST'])
def remove_project_logo():
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("UPDATE project SET logo_path = NULL")
        conn.commit()
        conn.close()
        flash('✅ لوگوی پروژه حذف شد.', 'success')
    except Exception as e:
        flash(f'خطا در حذف لوگوی پروژه: {str(e)}', 'error')
    return redirect(url_for('index'))

# ============================================================
# قسمت‌ها
# ============================================================
@app.route('/add_part', methods=['POST'])
def add_part():
    try:
        project_id = request.form.get('project_id')
        part_name = request.form.get('part_name').strip()
        if not part_name:
            flash('نام قسمت نمی‌تواند خالی باشد.', 'error')
            return redirect(url_for('index'))
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT id FROM project_parts WHERE project_id=? AND part_name=?", (project_id, part_name))
        if cur.fetchone():
            flash('این نام قسمت قبلاً ثبت شده است.', 'error')
            conn.close()
            return redirect(url_for('index'))
        cur.execute("INSERT INTO project_parts (project_id, part_name, created_at) VALUES (?, ?, ?)", (project_id, part_name, datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
        conn.commit()
        conn.close()
        flash('✅ قسمت/مرحله با موفقیت اضافه شد.', 'success')
        return redirect(url_for('index', active_tab='tab-project'))
    except Exception as e:
        flash(f'خطا: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/edit_part', methods=['POST'])
def edit_part():
    try:
        part_id = request.form.get('part_id')
        new_name = request.form.get('new_name', '').strip()
        if not new_name:
            flash('نام جدید نمی‌تواند خالی باشد.', 'error')
            return redirect(url_for('index', active_tab='tab-project'))
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT project_id FROM project_parts WHERE id=?", (part_id,))
        part = cur.fetchone()
        if not part:
            flash('قسمت مورد نظر یافت نشد.', 'error')
            conn.close()
            return redirect(url_for('index', active_tab='tab-project'))
        project_id = part['project_id']
        cur.execute("UPDATE project_parts SET part_name=? WHERE id=?", (new_name, part_id))
        conn.commit()
        cur.execute("SELECT project_name FROM project WHERE id=?", (project_id,))
        proj = cur.fetchone()
        conn.close()
        if proj:
            flash(f'✅ نام قسمت در پروژه "{proj["project_name"]}" به "{new_name}" تغییر یافت.', 'success')
        else:
            flash('✅ قسمت/مرحله ویرایش شد.', 'success')
        return redirect(url_for('index', active_tab='tab-project'))
    except Exception as e:
        flash(f'خطا: {str(e)}', 'error')
        return redirect(url_for('index', active_tab='tab-project'))

@app.route('/delete_part', methods=['POST'])
def delete_part():
    try:
        part_id = request.form.get('part_id')
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT part_name FROM project_parts WHERE id=?", (part_id,))
        part = cur.fetchone()
        if part and part['part_name'] == 'سایر':
            flash('قسمت پیش‌فرض "سایر" قابل حذف نیست.', 'error')
            conn.close()
            return redirect(url_for('index'))
        cur.execute("DELETE FROM project_parts WHERE id=?", (part_id,))
        conn.commit()
        conn.close()
        flash('✅ قسمت/مرحله حذف شد.', 'success')
    except Exception as e:
        flash(f'خطا: {str(e)}', 'error')
    return redirect(url_for('index'))

# ============================================================
# بازدید
# ============================================================
@app.route('/add_visit', methods=['POST'])
def add_visit():
    try:
        project_id = request.form.get('project_id')
        report_no_visit = request.form.get('report_no_visit')
        visit_date = request.form.get('visit_date')
        stage = request.form.get('stage')
        part_name = request.form.get('part_name')

        conn = get_db()
        cur = conn.cursor()
        
        cur.execute("SELECT id FROM defect WHERE project_id=? AND report_no=?", (project_id, report_no_visit))
        if cur.fetchone():
            flash('❌ این شماره گزارش قبلاً برای این پروژه ثبت شده است! لطفاً از شماره دیگری استفاده کنید.', 'error')
            conn.close()
            return redirect(url_for('index'))

        cur.execute("SELECT report_no, project_name, employer FROM project WHERE id=?", (project_id,))
        p = cur.fetchone()
        if not p:
            flash('پروژه یافت نشد', 'error')
            conn.close()
            return redirect(url_for('index'))
        miladi_date = visit_date
        try:
            parts = visit_date.split('/')
            if len(parts) == 3:
                shamsi = jdatetime.date(int(parts[0]), int(parts[1]), int(parts[2]))
                miladi_date = shamsi.togregorian().strftime("%Y-%m-%d")
        except: pass
        cur.execute("INSERT INTO defect (project_id, report_no, project_name, employer, visit_date, stage, is_active, part_name) VALUES (?, ?, ?, ?, ?, ?, ?, ?)", (project_id, report_no_visit, p['project_name'], p['employer'], miladi_date, stage, 1, part_name))
        inserted_id = cur.lastrowid
        conn.commit()
        conn.close()

        session['current_visit'] = {
            'defect_id': inserted_id,
            'report_no': report_no_visit,
            'project_name': p['project_name'],
            'part_name': part_name,
            'visit_date': miladi_date
        }
        session['defect_counter'] = 1

        flash('✅ بازدید با موفقیت ثبت شد. به تب انطباق منتقل شدید.', 'success')
        return redirect(url_for('index', active_tab='tab-defect'))
    except Exception as e:
        flash(f'خطا در ثبت بازدید: {str(e)}', 'error')
        return redirect(url_for('index'))

# ============================================================
# انطباق (ثبت جدید)
# ============================================================
@app.route('/add_defect', methods=['POST'])
def add_defect():
    try:
        defect_id = request.form.get('defect_id')
        title = request.form.get('title', '').strip()
        standard = request.form.get('standard', '').strip()
        description = request.form.get('description', '').strip()
        action = request.form.get('action', 'finalize')

        if not defect_id:
            flash('شناسه گزارش یافت نشد.', 'error')
            return redirect(url_for('index'))

        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT title, standard, description, image FROM defect WHERE id=?", (defect_id,))
        old = cur.fetchone()
        if not old:
            flash('گزارش یافت نشد', 'error')
            conn.close()
            return redirect(url_for('index'))
        
        img_paths = []
        files = request.files.getlist('image')
        for file in files:
            if file and file.filename:
                filename = secure_filename(f"defect_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{len(img_paths)}.jpg")
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                img_paths.append(filepath)
        img_path_str = "&&&".join(img_paths) if img_paths else ''
        
        sep = " ||| "
        nt = old['title'] + sep + title if old['title'] else title
        ns = old['standard'] + sep + standard if old['standard'] else standard
        nd = old['description'] + sep + description if old['description'] else description
        ni = old['image'] + sep + img_path_str if old['image'] else img_path_str
        
        cur.execute("UPDATE defect SET title=?, standard=?, description=?, image=? WHERE id=?", (nt, ns, nd, ni, defect_id))
        conn.commit()
        conn.close()

        if action == 'continue':
            session['defect_counter'] = session.get('defect_counter', 1) + 1
            flash(f'✅ مورد جدید با موفقیت ثبت شد. در حال ثبت مورد شماره {session["defect_counter"]} هستید.', 'success')
            return redirect(url_for('index', active_tab='tab-defect'))
        else:
            conn2 = get_db()
            cur2 = conn2.cursor()
            cur2.execute("""
                SELECT d.id as defect_id, d.report_no, p.id as project_id, p.project_name, d.part_name, d.stage 
                FROM defect d 
                JOIN project p ON d.project_id = p.id 
                WHERE d.id=?
            """, (defect_id,))
            report_data = cur2.fetchone()
            conn2.close()
            if report_data:
                session['current_report'] = dict(report_data)
            else:
                session.pop('current_report', None)
            session.pop('current_visit', None)
            session.pop('defect_counter', None)
            flash('✅ انطباق با موفقیت ثبت شد. به تب گزارش نهایی منتقل شدید.', 'success')
            return redirect(url_for('index', active_tab='tab-report'))
    except Exception as e:
        flash(f'خطا در ثبت انطباق: {str(e)}', 'error')
        return redirect(url_for('index'))

# ============================================================
# ویرایش گزارش
# ============================================================
@app.route('/edit_defect/<int:defect_id>', methods=['POST'])
def edit_defect(defect_id):
    session['edit_defect_id'] = defect_id
    session.pop('current_visit', None)
    flash('حالت ویرایش فعال شد. می‌توانید فیلدها را تغییر دهید.', 'info')
    return redirect(url_for('index', active_tab='tab-defect'))

@app.route('/update_defect/<int:defect_id>', methods=['POST'])
def update_defect(defect_id):
    try:
        title = request.form.get('title', '').strip()
        standard = request.form.get('standard', '').strip()
        description = request.form.get('description', '').strip()
        
        delete_list = request.form.getlist('delete_images[]')
        deleted = {}
        for item in delete_list:
            if item and ':' in item:
                g_idx, i_idx = item.split(':')
                try:
                    g_idx = int(g_idx)
                    i_idx = int(i_idx)
                    if g_idx not in deleted:
                        deleted[g_idx] = set()
                    deleted[g_idx].add(i_idx)
                except:
                    continue

        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT image FROM defect WHERE id=?", (defect_id,))
        old = cur.fetchone()
        old_images_str = old['image'] if old else ''

        groups = []
        if old_images_str:
            for group_str in old_images_str.split(' ||| '):
                imgs = [img.strip() for img in group_str.split('&&&') if img.strip()]
                if imgs:
                    groups.append(imgs)

        for g_idx, indices in deleted.items():
            if g_idx < len(groups):
                for i_idx in sorted(indices, reverse=True):
                    if i_idx < len(groups[g_idx]):
                        del groups[g_idx][i_idx]

        groups = [g for g in groups if g]

        for key in request.files:
            if key.startswith('new_images_group_'):
                try:
                    g_idx = int(key.split('_')[-1])
                except:
                    continue
                files = request.files.getlist(key)
                for file in files:
                    if file and file.filename:
                        filename = secure_filename(f"defect_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.jpg")
                        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                        file.save(filepath)
                        while len(groups) <= g_idx:
                            groups.append([])
                        groups[g_idx].append(filepath)

        new_images_str = ' ||| '.join(['&&&'.join(g) for g in groups if g])

        cur.execute("UPDATE defect SET title=?, standard=?, description=?, image=? WHERE id=?", 
                    (title, standard, description, new_images_str, defect_id))
        conn.commit()
        conn.close()

        session.pop('edit_defect_id', None)

        conn2 = get_db()
        cur2 = conn2.cursor()
        cur2.execute("""
            SELECT d.id as defect_id, d.report_no, p.id as project_id, p.project_name, d.part_name, d.stage 
            FROM defect d 
            JOIN project p ON d.project_id = p.id 
            WHERE d.id=?
        """, (defect_id,))
        report_data = cur2.fetchone()
        conn2.close()
        if report_data:
            session['current_report'] = dict(report_data)
            flash('✅ گزارش با موفقیت ویرایش شد. به تب گزارش نهایی منتقل شدید.', 'success')
        else:
            flash('⚠️ گزارش ویرایش شد اما اطلاعات گزارش یافت نشد.', 'warning')

        return redirect(url_for('index', active_tab='tab-report'))
    except Exception as e:
        flash(f'خطا در ویرایش گزارش: {str(e)}', 'error')
        return redirect(url_for('index', active_tab='tab-defect'))

# ============================================================
# حذف گزارش
# ============================================================
@app.route('/delete_defect/<int:defect_id>', methods=['POST'])
def delete_defect(defect_id):
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM defect WHERE id=?", (defect_id,))
        conn.commit()
        conn.close()
        flash('✅ گزارش با موفقیت حذف شد.', 'success')
    except Exception as e:
        flash(f'خطا در حذف گزارش: {str(e)}', 'error')
    return redirect(url_for('index', active_tab='tab-project'))

@app.route('/set_report/<int:defect_id>', methods=['POST'])
def set_report(defect_id):
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("""
            SELECT d.id as defect_id, d.report_no, p.id as project_id, p.project_name, d.part_name, d.stage 
            FROM defect d 
            JOIN project p ON d.project_id = p.id 
            WHERE d.id=?
        """, (defect_id,))
        row = cur.fetchone()
        conn.close()
        if row:
            session['current_report'] = dict(row)
            flash('✅ به تب گزارش نهایی منتقل شدید. می‌توانید گزارش را تولید کنید.', 'success')
        else:
            flash('❌ گزارش مورد نظر یافت نشد.', 'error')
    except Exception as e:
        flash(f'خطا در تنظیم گزارش: {str(e)}', 'error')
    return redirect(url_for('index', active_tab='tab-report'))

# ============================================================
# گزارش
# ============================================================
@app.route('/generate_report', methods=['POST'])
def generate_report():
    ok, msg = check_report_libs()
    if not ok:
        flash(f'⚠️ خطا در بارگذاری کتابخانه‌های گزارش:\n{msg}\n\nلطفاً دستورات زیر را در ترمینال اجرا کنید:\n1. pkg install libjpeg-turbo libpng libtiff libxml2 libxslt clang -y\n2. pip install reportlab python-docx', 'error')
        return redirect(url_for('index'))

    try:
        report_id = request.form.get('report_id')
        if not report_id:
            flash('شناسه گزارش ارسال نشد. لطفاً دوباره تلاش کنید.', 'error')
            return redirect(url_for('index', active_tab='tab-report'))

        header = request.form.get('header')
        result = request.form.get('result')
        pdf_enabled = 'pdf' in request.form
        word_enabled = 'word' in request.form
        
        conn = get_db()
        cur = conn.cursor()
        cur.execute("""
            SELECT d.*, 
                   p.project_name as proj_name, 
                   p.report_no as proj_report_no, 
                   p.address as proj_address,
                   p.logo_path as proj_logo,
                   e.fullname, e.position, e.license, e.phone, e.email, 
                   e.logo_path as user_logo,
                   e.signature_path, e.stamp_path
            FROM defect d
            JOIN project p ON d.project_id = p.id
            LEFT JOIN expert e ON e.id = (SELECT id FROM expert LIMIT 1)
            WHERE d.id=?
        """, (report_id,))
        data = cur.fetchone()
        conn.close()
        if not data:
            flash('اطلاعات گزارش یافت نشد', 'error')
            return redirect(url_for('index', active_tab='tab-report'))
        data = dict(data)

        project_report_no = data['proj_report_no']
        part_name = data['part_name'] or 'سایر'
        stage_display = data['stage'] or ''
        import re
        project_report_clean = re.sub(r'[\\/*?:"<>|]', '_', str(project_report_no))
        part_clean = re.sub(r'[\\/*?:"<>|]', '_', str(part_name))
        report_clean = re.sub(r'[\\/*?:"<>|]', '_', str(data['report_no']))
        
        part_folder = os.path.join(REPORTS_DIR, project_report_clean, part_clean)
        if not os.path.exists(part_folder):
            try: os.makedirs(part_folder)
            except: pass

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = f"{project_report_clean}_{report_clean}_{timestamp}"
        file_path_pdf = os.path.join(part_folder, f"{base_name}.pdf")
        file_path_docx = os.path.join(part_folder, f"{base_name}.docx")
        generated_files = []

        if pdf_enabled:
            try:
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.utils import ImageReader
                from reportlab.lib import colors
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont

                font_name = 'Helvetica'
                vazir_path = '/storage/emulated/0/Fonts/Vazir.ttf'
                try:
                    if os.path.exists(vazir_path):
                        pdfmetrics.registerFont(TTFont('CustomFont', vazir_path))
                        font_name = 'CustomFont'
                except: pass

                c = canvas.Canvas(file_path_pdf, pagesize=A4)
                width, height = A4; margin = 40; page_width = width - (2 * margin)
                c.setStrokeColor(colors.black); c.setLineWidth(1)
                c.rect(margin, margin, page_width, height - (2 * margin))

                y = height - margin - 15
                x_left = margin + 10; l_y = y; c.setFont(font_name, 9)
                if data.get('visit_date'):
                    c.drawString(x_left, l_y, fa(f"تاریخ بازدید: {convert_to_shamsi(data['visit_date'])}")); l_y -= 14
                if data.get('report_no'):
                    c.drawString(x_left, l_y, fa(f"شماره گزارش: {data['report_no']}")); l_y -= 14
                if data.get('proj_address'):
                    c.drawString(x_left, l_y, fa(f"آدرس: {data['proj_address']}")); l_y -= 14

                x_right = width - margin - 10; r_y = y
                user_logo = data.get('user_logo')
                if user_logo and os.path.exists(user_logo):
                    try: c.drawImage(ImageReader(user_logo), x_right-50, r_y-30, width=50, height=35, preserveAspectRatio=True); r_y -= 40
                    except: pass
                c.setFont(font_name, 9)
                if data.get('fullname'):
                    c.drawRightString(x_right, r_y, fa(f"نام کارشناس: {data['fullname']}")); r_y -= 14
                if data.get('position'):
                    c.drawRightString(x_right, r_y, fa(f"سمت: {data['position']}")); r_y -= 14

                x_center = width / 2; c_y = y
                proj_logo = data.get('proj_logo')
                if proj_logo and os.path.exists(proj_logo):
                    try: c.drawImage(ImageReader(proj_logo), x_center-30, c_y-35, width=60, height=40, preserveAspectRatio=True); c_y -= 45
                    except: pass
                c.setFont(font_name, 10)
                if data.get('proj_name'):
                    tw = c.stringWidth(fa(data['proj_name']), font_name, 10)
                    c.drawString(x_center - tw/2, c_y, fa(data['proj_name'])); c_y -= 16
                c.setFont(font_name, 9)
                if part_name:
                    tw = c.stringWidth(fa(f"قسمت/مرحله: {part_name}"), font_name, 9)
                    c.drawString(x_center - tw/2, c_y, fa(f"قسمت/مرحله: {part_name}")); c_y -= 14
                if stage_display:
                    tw = c.stringWidth(fa(stage_display), font_name, 9)
                    c.drawString(x_center - tw/2, c_y, fa(stage_display)); c_y -= 14

                line_y = y - 85; c.setStrokeColor(colors.black); c.setLineWidth(0.5)
                c.line(x_center - page_width/6, line_y, x_center + page_width/6, line_y); y = line_y - 15

                if header.strip():
                    c.setFont(font_name, 11); lines = split_text_to_lines(header, c, font_name, page_width - 20, 11); y -= 10
                    for line in lines: c.drawRightString(width - margin - 10, y, fa(line)); y -= 20
                    y -= 10

                titles = data.get('title', '').split(" ||| ") if data.get('title') else []
                standards = data.get('standard', '').split(" ||| ") if data.get('standard') else []
                descriptions = data.get('description', '').split(" ||| ") if data.get('description') else []
                images = data.get('image', '').split(" ||| ") if data.get('image') else []

                for idx, title in enumerate(titles):
                    img_paths = images[idx].split("&&&") if idx < len(images) and images[idx] else []
                    c.setFont(font_name, 10)
                    text_parts = [f"{idx+1} - عنوان: {title}"]
                    if idx < len(standards) and standards[idx]: text_parts.append(f"استاندارد: {standards[idx]}")
                    if idx < len(descriptions) and descriptions[idx]: text_parts.append(f"شرح: {descriptions[idx]}")

                    for part in text_parts:
                        lines = split_text_to_lines(part, c, font_name, page_width - 20, 10)
                        for line in lines:
                            if y - 18 < margin + 20:
                                c.showPage(); c.rect(margin, margin, page_width, height - (2 * margin)); y = height - margin - 15
                            c.drawRightString(width - margin - 10, y, fa(line)); y -= 18

                    if img_paths:
                        img_height = 80; y -= 5
                        valid_imgs = [p for p in img_paths if os.path.exists(p.strip())]
                        if valid_imgs:
                            cols = min(4, len(valid_imgs)); rows = (len(valid_imgs) + cols - 1) // cols
                            total_h = rows * (img_height + 10)
                            if y - total_h < margin + 20:
                                c.showPage(); c.rect(margin, margin, page_width, height - (2 * margin)); y = height - margin - 15
                            start_y = y - 10
                            for i, path in enumerate(valid_imgs):
                                try:
                                    img = ImageReader(path.strip())
                                    col = i % cols; row = i // cols
                                    pos_x = margin + 10 + col * ((page_width - 20) / cols)
                                    pos_y = start_y - row * (img_height + 10)
                                    c.drawImage(img, pos_x, pos_y - img_height, width=80, height=img_height, preserveAspectRatio=True)
                                    c.setFont(font_name, 8)
                                    num_text = f"({i+1})"
                                    tw = c.stringWidth(num_text, font_name, 8)
                                    c.drawString(pos_x + 80 + 5, pos_y - img_height/2 - 4, num_text)
                                except: c.drawString(pos_x + 40, pos_y - (img_height/2), fa("[تصویر موجود نیست]"))
                            y = start_y - (rows * (img_height + 10)) - 10

                if result.strip():
                    y -= 10; c.setFont(font_name, 11)
                    c.drawRightString(width - margin - 10, y, fa("نتیجه بازدید:")); y -= 15
                    c.setFont(font_name, 10); lines = split_text_to_lines(result, c, font_name, page_width - 20, 10)
                    for line in lines:
                        if y - 16 < margin: c.showPage(); c.rect(margin, margin, page_width, height - (2 * margin)); y = height - 40
                        c.drawRightString(width - margin - 10, y, fa(line)); y -= 16

                cc_list = request.form.getlist('cc[]')
                if cc_list and any(cc.strip() for cc in cc_list):
                    y -= 10
                    c.setFont(font_name, 8)
                    c.drawRightString(width - margin - 10, y, fa("رونوشت:"));
                    y -= 15
                    for cc in cc_list:
                        if cc.strip():
                            c.drawRightString(width - margin - 10, y, fa(f"• {cc.strip()}"))
                            y -= 12

                sig_path = data.get('signature_path')
                try:
                    if sig_path and os.path.exists(sig_path): c.drawImage(ImageReader(sig_path), margin + 10, y - 40, width=80, height=50)
                except: c.drawString(margin + 10, y - 20, fa("[امضا]"))
                stamp_path = data.get('stamp_path')
                try:
                    if stamp_path and os.path.exists(stamp_path): c.drawImage(ImageReader(stamp_path), margin + 100, y - 40, width=80, height=50)
                except: c.drawString(margin + 100, y - 20, fa("[مهر]"))

                c.save()
                generated_files.append(f"PDF: {base_name}.pdf")
            except Exception as e:
                flash(f'❌ خطا در تولید PDF:\n{str(e)}', 'error')
                return redirect(url_for('index'))

        if word_enabled:
            try:
                from docx import Document
                from docx.shared import Pt
                doc = Document()
                doc.add_heading(f"گزارش MHB - {data['report_no']}", 0)
                doc.add_paragraph(f"پروژه: {data['proj_name']}")
                doc.add_paragraph(f"قسمت/مرحله: {part_name}")
                doc.add_paragraph(f"نوبت: {stage_display}")
                doc.add_paragraph(f"نتیجه: {result}")
                
                cc_list = request.form.getlist('cc[]')
                if cc_list and any(cc.strip() for cc in cc_list):
                    p = doc.add_paragraph("رونوشت:")
                    p.runs[0].font.size = Pt(8)
                    for cc in cc_list:
                        if cc.strip():
                            p = doc.add_paragraph(f"• {cc.strip()}")
                            p.runs[0].font.size = Pt(8)
                            
                doc.save(file_path_docx)
                generated_files.append(f"Word: {base_name}.docx")
            except Exception as e:
                flash(f'❌ خطا در تولید Word:\n{str(e)}', 'error')

        if generated_files:
            folder_link = f'<a href="/download_file?path={file_path_pdf}" target="_blank" style="color: #3498db; font-weight: bold;">⬇️ دانلود مستقیم فایل PDF</a>'
            flash(f'✅ گزارش تولید شد! فایل‌ها: {", ".join(generated_files)}\n\n🔗 {folder_link}', 'success')
            session.pop('current_report', None)
            session['report_just_generated'] = True
        else:
            flash('⚠️ هیچ فایلی تولید نشد.', 'warning')
    except Exception as e:
        flash(f'❌ خطا: {str(e)}', 'error')
    return redirect(url_for('index', active_tab='tab-report'))

@app.route('/download_file')
def download_file():
    path = request.args.get('path')
    if path and os.path.exists(path):
        return send_file(path, as_attachment=True)
    flash('فایل یافت نشد.', 'error')
    return redirect(url_for('index'))

@app.route('/open_folder')
def open_folder():
    path = request.args.get('path')
    if path and os.path.exists(path):
        flash(f'📂 مسیر پوشه گزارش: {path}', 'info')
    else:
        flash('مسیر پوشه یافت نشد', 'error')
    return redirect(url_for('index'))

# ============================================================
# اجرای برنامه در اندروید (بخش Kivy کامنت شده تا در ترموکس اجرا شود)
# ============================================================
import threading
import time

def run_flask_server():
    app.run(host='127.0.0.1', port=5000, debug=True, use_reloader=False, threaded=True)

if __name__ == '__main__':
     
     from kivymd.app import MDApp
     from kivymd.uix.webview import MDWebView
     from kivy.uix.boxlayout import BoxLayout
     class MainApp(MDApp):
         def build(self):
             self.server_thread = threading.Thread(target=run_flask_server, daemon=True)
             self.server_thread.start()
             time.sleep(1)
             layout = BoxLayout()
             webview = MDWebView(url="http://127.0.0.1:5000", enable_javascript=True)
             layout.add_widget(webview)
             return layout
     MainApp().run()
    
    
